"""Build a Word handoff packet for external scientific review.

Author: RMTGuard contributors
Date: 2026-05-12
Purpose: Convert the current evidence-freeze manuscript text, figure legends,
and journal-route gap assessment into a reviewer-friendly DOCX packet.
Data sources: local Markdown/TSV artifacts generated from the RMTGuard
benchmark and submission-control pipeline.
Method notes: Uses python-docx for reproducible document generation.
"""

from __future__ import annotations

import csv
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "doc" / "RMTGuard_external_review_packet_2026-05-12.docx"

INPUTS = {
    "results": ROOT / "manuscript" / "results_freeze_aligned_draft.md",
    "legends": ROOT / "manuscript" / "figure_legends_freeze_aligned.md",
    "dashboard": ROOT / "docs" / "high_impact_submission_dashboard.md",
    "gap": ROOT / "docs" / "jif20_50_gap_assessment.md",
    "freeze": ROOT / "docs" / "current_evidence_freeze_2026-05-12.md",
    "reporting": ROOT / "docs" / "nature_reporting_summary_draft.md",
    "figure_audit": ROOT / "docs" / "figure_caption_source_audit.md",
    "version_audit": ROOT / "docs" / "post_release_version_coverage_audit.md",
    "v011_preflight": ROOT / "docs" / "v0_1_1_release_preflight.md",
    "author_declarations": ROOT / "docs" / "author_declaration_confirmation_packet.md",
    "manual": ROOT / "docs" / "manual_next_actions_20_50.md",
    "audit": ROOT / "results" / "submission" / "freeze_aligned_text_audit.tsv",
    "gantt": ROOT / "figures" / "project_management" / "rmtguard_project_gantt.png",
}


def read_text(path: Path) -> str:
    if not path.exists():
        return f"[MISSING: {path.as_posix()}]"
    return path.read_text(encoding="utf-8")


def add_hyperlink_text(paragraph, label: str, url: str) -> None:
    """Add URL text as plain readable text; avoids fragile relationship XML."""
    paragraph.add_run(label).bold = True
    paragraph.add_run(f": {url}")


def clean_inline(text: str) -> str:
    """Remove lightweight Markdown markers that look unprofessional in Word."""
    return text.replace("`", "")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(18)
    for style_name, size in [("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 10.5)]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.bold = True


def add_key_value_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Current value"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_markdown_section(document: Document, title: str, text: str, max_lines: int | None = None) -> None:
    document.add_heading(title, level=1)
    lines = text.splitlines()
    if max_lines is not None:
        lines = lines[:max_lines]
    in_code = False
    pending_table: list[list[str]] = []

    def flush_table() -> None:
        nonlocal pending_table
        if not pending_table:
            return
        table = document.add_table(rows=1, cols=len(pending_table[0]))
        table.style = "Table Grid"
        for idx, value in enumerate(pending_table[0]):
            table.rows[0].cells[idx].text = value
        for row in pending_table[1:]:
            cells = table.add_row().cells
            for idx, value in enumerate(row[: len(cells)]):
                cells[idx].text = value
        pending_table = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if line.startswith("```"):
            flush_table()
            in_code = not in_code
            continue
        if not line:
            flush_table()
            continue
        if line.startswith("|") and line.endswith("|"):
            cells = [clean_inline(cell.strip(" `")) for cell in line.strip("|").split("|")]
            if cells and all(set(cell) <= {"-", ":"} for cell in cells):
                continue
            pending_table.append(cells)
            continue
        flush_table()
        if in_code:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
        elif line.startswith("### "):
            document.add_heading(clean_inline(line[4:]), level=3)
        elif line.startswith("## "):
            document.add_heading(clean_inline(line[3:]), level=2)
        elif line.startswith("# "):
            document.add_heading(clean_inline(line[2:]), level=2)
        elif line.startswith("- "):
            document.add_paragraph(clean_inline(line[2:]), style="List Bullet")
        elif line and line[0].isdigit() and ". " in line[:5]:
            document.add_paragraph(clean_inline(line.split(". ", 1)[1]), style="List Number")
        else:
            document.add_paragraph(clean_inline(line))
    flush_table()


def add_audit_table(document: Document, audit_path: Path) -> None:
    document.add_heading("Machine-readable claim audit", level=1)
    if not audit_path.exists():
        document.add_paragraph(f"Missing audit file: {audit_path.as_posix()}")
        return
    with audit_path.open("r", encoding="utf-8", newline="") as handle:
        rows = list(csv.DictReader(handle, delimiter="\t"))
    if not rows:
        document.add_paragraph("Audit table is empty.")
        return
    cols = ["section", "claim_boundary", "evidence_source", "status"]
    table = document.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for idx, col in enumerate(cols):
        table.rows[0].cells[idx].text = col
    for row in rows:
        cells = table.add_row().cells
        for idx, col in enumerate(cols):
            cells[idx].text = row.get(col, "")
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_external_review_questions(document: Document) -> None:
    document.add_heading("Questions for external reviewers", level=1)
    questions = [
        "Is the central claim best framed as random-matrix callability/noise control rather than broad stability superiority?",
        "Which benchmark result would most likely trigger reviewer skepticism, and what extra analysis would reduce that risk?",
        "Does the no-call decision map look like a useful methodological feature, or does it read as underperformance?",
        "Is the PDAC/TME public-data application strong enough for a main figure, or should it move to supplementary material?",
        "For journal routing, should the next live target remain Nature Methods presubmission, or should the manuscript be reframed directly for Genome Biology?",
    ]
    for question in questions:
        document.add_paragraph(question, style="List Number")


def build_document() -> None:
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RMTGuard External Review Packet")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Evidence-freeze aligned manuscript materials for pre-review").italic = True
    document.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    add_key_value_table(
        document,
        [
            ("Current route", "Nature Methods remains a stretch target; Genome Biology is the realistic fallback if the high-impact gate is not recovered."),
            ("Acceptance guarantee", "Impossible. The document is for risk reduction and external critique, not a promise of publication."),
            ("Public code", "Repository and DOI metadata are present for the archived release: https://github.com/healthgreat/rmtguard ; DOI 10.5281/zenodo.20012350."),
            ("Core claim boundary", "Diagnostic random-matrix callability and noise-control reporting; not broad superiority over every fixed-PC/elbow workflow."),
            ("Manual blocker", "Corresponding-author acknowledgement for bounded Figure 4 wording and final author declarations remain author-owned."),
        ],
    )

    paragraph = document.add_paragraph()
    add_hyperlink_text(paragraph, "GitHub repository", "https://github.com/healthgreat/rmtguard")
    paragraph = document.add_paragraph()
    add_hyperlink_text(paragraph, "Zenodo DOI", "https://doi.org/10.5281/zenodo.20012350")

    if INPUTS["gantt"].exists():
        document.add_heading("Project progress Gantt snapshot", level=1)
        document.add_picture(str(INPUTS["gantt"]), width=Inches(5.6))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_section(WD_SECTION.NEW_PAGE)
    add_markdown_section(document, "Current evidence freeze", read_text(INPUTS["freeze"]), max_lines=43)
    add_markdown_section(document, "High-impact submission dashboard", read_text(INPUTS["dashboard"]), max_lines=130)
    add_markdown_section(document, "20-50 JIF gap assessment", read_text(INPUTS["gap"]), max_lines=140)
    add_markdown_section(document, "Figure-caption-source audit", read_text(INPUTS["figure_audit"]), max_lines=120)
    add_markdown_section(document, "Post-release version coverage audit", read_text(INPUTS["version_audit"]), max_lines=120)
    add_markdown_section(document, "v0.1.1 release preflight", read_text(INPUTS["v011_preflight"]), max_lines=110)
    add_markdown_section(document, "Author declaration confirmation packet", read_text(INPUTS["author_declarations"]), max_lines=130)
    add_markdown_section(document, "Nature reporting-summary worksheet", read_text(INPUTS["reporting"]), max_lines=150)
    add_markdown_section(document, "Freeze-aligned Results draft", read_text(INPUTS["results"]))
    add_markdown_section(document, "Freeze-aligned figure legends", read_text(INPUTS["legends"]))
    add_audit_table(document, INPUTS["audit"])
    add_external_review_questions(document)
    add_markdown_section(document, "Manual author-owned tasks", read_text(INPUTS["manual"]), max_lines=130)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = OUTPUT.with_suffix(".docx.tmp")
    document.save(tmp_path)
    if OUTPUT.exists():
        OUTPUT.unlink()
    tmp_path.replace(OUTPUT)
    print(f"wrote {OUTPUT.relative_to(ROOT).as_posix()}")


if __name__ == "__main__":
    build_document()
