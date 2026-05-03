#!/usr/bin/env python
"""Build real-data ablation figure and supplemental table assets.

Author: RMTGuard development team
Date: 2026-05-03
Purpose: Convert the four-dataset `subsample80_pilot10` real-data annotation
ablation screen into reviewer-facing source data, a forest plot, a
supplemental table, and an evidence note.
Data source: results/ablation/realdata_ablation_annotation_detail.tsv.
Method notes: This script summarizes existing pilot results only. It does not
change algorithm outputs, rerun benchmarks, or upgrade pilot evidence to a
final manuscript-grade claim.
"""

from __future__ import annotations

import csv
import math
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
DETAIL = ROOT / "results" / "ablation" / "realdata_ablation_annotation_detail.tsv"
SOURCE_DATA = (
    ROOT
    / "results"
    / "figures"
    / "source_data"
    / "figure5_realdata_ablation_delta_summary.tsv"
)
SUPP_TABLE_TSV = (
    ROOT / "results" / "tables" / "manuscript" / "supplemental_realdata_ablation_table.tsv"
)
SUPP_TABLE_DOCX = (
    ROOT / "results" / "tables" / "manuscript" / "supplemental_realdata_ablation_table.docx"
)
FIG_DIR = ROOT / "figures" / "manuscript"
FIG_PNG = FIG_DIR / "figure5_realdata_ablation_forest.png"
FIG_PDF = FIG_DIR / "figure5_realdata_ablation_forest.pdf"
FIG_TIFF = FIG_DIR / "figure5_realdata_ablation_forest.tiff"
FIG_MANIFEST = FIG_DIR / "realdata_ablation_figure_manifest.tsv"
TABLE_MANIFEST = (
    ROOT / "results" / "tables" / "manuscript" / "realdata_ablation_table_manifest.tsv"
)
DOC = ROOT / "docs" / "realdata_ablation_figure_table.md"

RUN_LABEL = "subsample80_pilot10"
REFERENCE_ABLATION = "default_v3_3"

DATASET_ORDER = [
    "baron_pancreas",
    "kang_ifnb_pbmc",
    "pbmc68k_zheng2017",
    "pdac_gse263733",
]
ABLATION_ORDER = [
    "embedding_strict_signal",
    "force_min_embedding_pcs_10",
    "rare_state_guard_off",
    "batch_residualized",
]

DATASET_LABELS = {
    "baron_pancreas": "Baron pancreas",
    "kang_ifnb_pbmc": "Kang IFN-beta PBMC",
    "pbmc68k_zheng2017": "PBMC68k/Zheng 2017",
    "pdac_gse263733": "PDAC GSE263733",
}

VARIANT_LABELS = {
    "default_v3_3": "Default RMTGuard",
    "embedding_strict_signal": "Strict signal embedding",
    "force_min_embedding_pcs_10": "Forced min 10 PCs",
    "rare_state_guard_off": "Rare-state guard off",
    "batch_residualized": "Batch residualized",
}

VARIANT_COLORS = {
    "embedding_strict_signal": "#4C78A8",
    "force_min_embedding_pcs_10": "#E45756",
    "rare_state_guard_off": "#72B7B2",
    "batch_residualized": "#54A24B",
}


def _rel(path: Path) -> str:
    try:
        return str(path.resolve().relative_to(ROOT.resolve())).replace("\\", "/")
    except ValueError:
        return str(path)


def _atomic_write_tsv(df: pd.DataFrame, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    df.to_csv(tmp, sep="\t", index=False, quoting=csv.QUOTE_MINIMAL)
    tmp.replace(path)


def _atomic_write_text(text: str, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(text.rstrip() + "\n", encoding="utf-8")
    tmp.replace(path)


def _atomic_save_figure(fig: plt.Figure, path: Path, dpi: int = 300) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    fig.savefig(
        tmp,
        format=path.suffix.lstrip("."),
        dpi=dpi,
        bbox_inches="tight",
        facecolor="white",
        metadata={"Creator": "RMTGuard build_realdata_ablation_assets.py"}
        if path.suffix.lower() != ".tiff"
        else None,
    )
    tmp.replace(path)


def _finite(values: pd.Series | np.ndarray | list[float]) -> np.ndarray:
    arr = pd.to_numeric(pd.Series(values), errors="coerce").to_numpy(dtype=float)
    return arr[np.isfinite(arr)]


def _mean_ci(values: pd.Series | np.ndarray | list[float]) -> tuple[float, float, float]:
    arr = _finite(values)
    if len(arr) == 0:
        return math.nan, math.nan, math.nan
    mean = float(np.mean(arr))
    if len(arr) == 1:
        return mean, mean, mean
    half_width = 1.96 * float(np.std(arr, ddof=1)) / math.sqrt(len(arr))
    return mean, mean - half_width, mean + half_width


def _format_ci(mean: float, low: float, high: float) -> str:
    if not np.isfinite(mean):
        return "NA"
    return f"{mean:.3f} ({low:.3f} to {high:.3f})"


def _interpretation(dataset_id: str, ablation_id: str, no_call_rate: float) -> str:
    if ablation_id == REFERENCE_ABLATION:
        if dataset_id == "pbmc68k_zheng2017":
            return "reference diagnostic no-call stress context; not a positive discovery claim"
        return "reference"
    if dataset_id == "pbmc68k_zheng2017" and ablation_id == "force_min_embedding_pcs_10":
        return (
            "forced clustering weakly increases annotation ARI but removes the "
            "no-call guard; negative control only"
        )
    if dataset_id == "pbmc68k_zheng2017":
        return "diagnostic stress/no-call context; not a rescue"
    if ablation_id == "batch_residualized" and dataset_id == "baron_pancreas":
        return "batch-aware signal: higher label ARI and lower batch ARI"
    if ablation_id == "batch_residualized" and dataset_id == "pdac_gse263733":
        return "external batch-aware signal: lower batch ARI with similar label ARI"
    if ablation_id == "batch_residualized" and dataset_id == "kang_ifnb_pbmc":
        return "lower batch ARI but no label ARI gain over default"
    if ablation_id == "force_min_embedding_pcs_10":
        return "forced-PC sensitivity control; do not promote as default"
    if ablation_id == "embedding_strict_signal":
        return "strict-signal embedding ablation"
    if ablation_id == "rare_state_guard_off":
        return "rare-state guard ablation"
    if no_call_rate > 0:
        return "contains diagnostic no-call repeats"
    return "component sensitivity check"


def build_summary() -> pd.DataFrame:
    if not DETAIL.exists():
        raise FileNotFoundError(f"Missing input table: {_rel(DETAIL)}")
    detail = pd.read_csv(DETAIL, sep="\t")
    required = {
        "run_label",
        "dataset_id",
        "ablation_id",
        "repeat",
        "variant_label",
        "analysis_status",
        "label_ari",
        "batch_ari",
        "cluster_n",
    }
    missing = sorted(required - set(detail.columns))
    if missing:
        raise ValueError(f"Input table is missing required columns: {missing}")
    sub = detail[detail["run_label"] == RUN_LABEL].copy()
    if sub.empty:
        raise ValueError(f"No rows found for run_label={RUN_LABEL!r}")

    ref = sub[sub["ablation_id"] == REFERENCE_ABLATION][
        ["dataset_id", "repeat", "label_ari", "batch_ari"]
    ].rename(columns={"label_ari": "ref_label_ari", "batch_ari": "ref_batch_ari"})
    rows: list[dict[str, object]] = []
    for (dataset_id, ablation_id), group in sub.groupby(["dataset_id", "ablation_id"]):
        merged = group.merge(ref, on=["dataset_id", "repeat"], how="left")
        label_delta = pd.to_numeric(merged["label_ari"], errors="coerce") - pd.to_numeric(
            merged["ref_label_ari"], errors="coerce"
        )
        batch_delta = pd.to_numeric(merged["batch_ari"], errors="coerce") - pd.to_numeric(
            merged["ref_batch_ari"], errors="coerce"
        )
        label_mean, label_low, label_high = _mean_ci(group["label_ari"])
        batch_mean, batch_low, batch_high = _mean_ci(group["batch_ari"])
        label_delta_mean, label_delta_low, label_delta_high = _mean_ci(label_delta)
        batch_delta_mean, batch_delta_low, batch_delta_high = _mean_ci(batch_delta)
        no_call_rate = float(
            (group["analysis_status"].astype(str) == "diagnostic_no_call").mean()
        )
        rows.append(
            {
                "run_label": RUN_LABEL,
                "dataset_id": dataset_id,
                "dataset_label": DATASET_LABELS.get(
                    dataset_id, dataset_id.replace("_", " ")
                ),
                "dataset_order": DATASET_ORDER.index(dataset_id)
                if dataset_id in DATASET_ORDER
                else 999,
                "ablation_id": ablation_id,
                "variant_label": VARIANT_LABELS.get(
                    ablation_id,
                    str(group["variant_label"].iloc[0]).replace(" v3.3", ""),
                ),
                "variant_order": (
                    -1
                    if ablation_id == REFERENCE_ABLATION
                    else ABLATION_ORDER.index(ablation_id)
                    if ablation_id in ABLATION_ORDER
                    else 999
                ),
                "n_repeats": int(group["repeat"].nunique()),
                "mean_label_ari": label_mean,
                "label_ari_ci95_low": label_low,
                "label_ari_ci95_high": label_high,
                "label_ari_delta_vs_default_mean": label_delta_mean,
                "label_ari_delta_vs_default_ci95_low": label_delta_low,
                "label_ari_delta_vs_default_ci95_high": label_delta_high,
                "paired_n_label_delta": int(len(_finite(label_delta))),
                "mean_batch_ari": batch_mean,
                "batch_ari_ci95_low": batch_low,
                "batch_ari_ci95_high": batch_high,
                "batch_ari_delta_vs_default_mean": batch_delta_mean,
                "batch_ari_delta_vs_default_ci95_low": batch_delta_low,
                "batch_ari_delta_vs_default_ci95_high": batch_delta_high,
                "paired_n_batch_delta": int(len(_finite(batch_delta))),
                "no_call_rate": no_call_rate,
                "mean_cluster_n": float(
                    pd.to_numeric(group["cluster_n"], errors="coerce").mean()
                ),
                "interpretation": _interpretation(
                    dataset_id, ablation_id, no_call_rate
                ),
            }
        )
    out = pd.DataFrame(rows)
    out = out.sort_values(["dataset_order", "variant_order", "ablation_id"]).reset_index(
        drop=True
    )
    return out


def render_forest(summary: pd.DataFrame) -> None:
    plot_df = summary[summary["ablation_id"] != REFERENCE_ABLATION].copy()
    plot_df = plot_df.sort_values(
        ["dataset_order", "variant_order", "ablation_id"], ascending=[True, True, True]
    ).reset_index(drop=True)
    plot_df["row_label"] = (
        plot_df["dataset_label"].astype(str)
        + " | "
        + plot_df["variant_label"].astype(str)
    )

    plt.rcParams.update(
        {
            "font.family": "Arial",
            "font.size": 8.0,
            "axes.titlesize": 9.0,
            "axes.labelsize": 8.0,
            "xtick.labelsize": 7.0,
            "ytick.labelsize": 7.0,
            "pdf.fonttype": 42,
            "ps.fonttype": 42,
        }
    )
    fig, axes = plt.subplots(
        1,
        2,
        figsize=(11.0, max(6.8, 0.32 * len(plot_df) + 1.8)),
        gridspec_kw={"width_ratios": [1.15, 1.0], "wspace": 0.35},
    )
    fig.subplots_adjust(left=0.31, right=0.98, top=0.93, bottom=0.16, wspace=0.35)

    y = np.arange(len(plot_df))[::-1]
    for ax, metric, low, high, title, xlabel in [
        (
            axes[0],
            "label_ari_delta_vs_default_mean",
            "label_ari_delta_vs_default_ci95_low",
            "label_ari_delta_vs_default_ci95_high",
            "A  Annotation ARI delta",
            "Delta vs default RMTGuard",
        ),
        (
            axes[1],
            "batch_ari_delta_vs_default_mean",
            "batch_ari_delta_vs_default_ci95_low",
            "batch_ari_delta_vs_default_ci95_high",
            "B  Batch ARI delta",
            "Delta vs default RMTGuard (lower is better)",
        ),
    ]:
        ax.axvline(0.0, color="#2F3437", linewidth=0.9)
        ax.grid(axis="x", color="#E6E8EB", linewidth=0.7)
        ax.set_axisbelow(True)
        for spine in ["top", "right"]:
            ax.spines[spine].set_visible(False)
        for i, row in plot_df.iterrows():
            mean = row[metric]
            if not np.isfinite(mean):
                continue
            xerr = np.array(
                [[mean - row[low]], [row[high] - mean]], dtype=float
            )
            color = VARIANT_COLORS.get(str(row["ablation_id"]), "#6B7280")
            ax.errorbar(
                mean,
                y[i],
                xerr=xerr,
                fmt="o",
                markersize=4.0,
                color=color,
                ecolor=color,
                elinewidth=1.0,
                capsize=2.2,
                markeredgecolor="white",
                markeredgewidth=0.5,
            )
        ax.set_title(title, loc="left", fontweight="bold")
        ax.set_xlabel(xlabel)
        ax.set_ylim(-0.8, len(plot_df) - 0.2)

    axes[0].set_yticks(y)
    axes[0].set_yticklabels(plot_df["row_label"])
    axes[1].set_yticks(y)
    axes[1].set_yticklabels([])
    axes[0].set_xlim(-0.12, 0.12)
    axes[1].set_xlim(-0.05, 0.05)

    handles = []
    labels = []
    for ablation_id in ABLATION_ORDER:
        handles.append(
            plt.Line2D(
                [0],
                [0],
                marker="o",
                color="none",
                markerfacecolor=VARIANT_COLORS[ablation_id],
                markeredgecolor="white",
                markersize=6,
            )
        )
        labels.append(VARIANT_LABELS[ablation_id])
    fig.legend(
        handles,
        labels,
        loc="lower center",
        ncol=4,
        frameon=False,
        bbox_to_anchor=(0.57, 0.055),
    )
    fig.text(
        0.01,
        0.018,
        "Pilot evidence only: 80% subsampling, 10 repeats per dataset/variant. "
        "PBMC68k is retained as a diagnostic no-call stress context.",
        fontsize=7.2,
        color="#4B5563",
    )
    _atomic_save_figure(fig, FIG_PNG)
    _atomic_save_figure(fig, FIG_PDF)
    _atomic_save_figure(fig, FIG_TIFF)
    plt.close(fig)


def build_supplemental_table(summary: pd.DataFrame) -> pd.DataFrame:
    table = summary.copy()
    table["label_ari_95ci"] = [
        _format_ci(row.mean_label_ari, row.label_ari_ci95_low, row.label_ari_ci95_high)
        for row in table.itertuples(index=False)
    ]
    table["delta_label_ari_95ci"] = [
        _format_ci(
            row.label_ari_delta_vs_default_mean,
            row.label_ari_delta_vs_default_ci95_low,
            row.label_ari_delta_vs_default_ci95_high,
        )
        for row in table.itertuples(index=False)
    ]
    table["batch_ari_95ci"] = [
        _format_ci(row.mean_batch_ari, row.batch_ari_ci95_low, row.batch_ari_ci95_high)
        for row in table.itertuples(index=False)
    ]
    table["delta_batch_ari_95ci"] = [
        _format_ci(
            row.batch_ari_delta_vs_default_mean,
            row.batch_ari_delta_vs_default_ci95_low,
            row.batch_ari_delta_vs_default_ci95_high,
        )
        for row in table.itertuples(index=False)
    ]
    keep = [
        "dataset_id",
        "dataset_label",
        "ablation_id",
        "variant_label",
        "n_repeats",
        "label_ari_95ci",
        "delta_label_ari_95ci",
        "batch_ari_95ci",
        "delta_batch_ari_95ci",
        "no_call_rate",
        "mean_cluster_n",
        "interpretation",
    ]
    return table[keep]


def write_docx(table: pd.DataFrame) -> None:
    SUPP_TABLE_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    title = doc.add_heading("Supplemental Table: Real-data ablation pilot", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note = doc.add_paragraph(
        "The table summarizes the existing 80% subsampling, 10-repeat pilot. "
        "It is not a final manuscript-grade 20-50 repeat analysis."
    )
    note.runs[0].font.size = Pt(9)
    columns = [
        "Dataset",
        "Variant",
        "n",
        "Label ARI (95% CI)",
        "Delta label ARI",
        "Batch ARI (95% CI)",
        "Delta batch ARI",
        "No-call",
        "Clusters",
        "Interpretation",
    ]
    doc_table = doc.add_table(rows=1, cols=len(columns))
    doc_table.style = "Table Grid"
    for cell, label in zip(doc_table.rows[0].cells, columns):
        cell.text = label
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(7)
    for row in table.itertuples(index=False):
        cells = doc_table.add_row().cells
        values = [
            row.dataset_label,
            row.variant_label,
            str(row.n_repeats),
            row.label_ari_95ci,
            row.delta_label_ari_95ci,
            row.batch_ari_95ci,
            row.delta_batch_ari_95ci,
            f"{row.no_call_rate:.2f}",
            f"{row.mean_cluster_n:.1f}",
            row.interpretation,
        ]
        for cell, value in zip(cells, values):
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(6.5)
    tmp = SUPP_TABLE_DOCX.with_suffix(SUPP_TABLE_DOCX.suffix + ".tmp")
    doc.save(tmp)
    tmp.replace(SUPP_TABLE_DOCX)


def write_manifests() -> None:
    fig_manifest = pd.DataFrame(
        [
            {
                "figure_id": "Figure 5 supplement: real-data ablation forest",
                "png_path": _rel(FIG_PNG),
                "pdf_path": _rel(FIG_PDF),
                "tiff_path": _rel(FIG_TIFF),
                "source_data": _rel(SOURCE_DATA),
                "status": "pilot_asset_generated",
            }
        ]
    )
    table_manifest = pd.DataFrame(
        [
            {
                "artifact": "supplemental_realdata_ablation_table_tsv",
                "path": _rel(SUPP_TABLE_TSV),
                "role": "supplemental_table_source",
            },
            {
                "artifact": "supplemental_realdata_ablation_table_docx",
                "path": _rel(SUPP_TABLE_DOCX),
                "role": "supplemental_table_word",
            },
        ]
    )
    _atomic_write_tsv(fig_manifest, FIG_MANIFEST)
    _atomic_write_tsv(table_manifest, TABLE_MANIFEST)


def build_doc(summary: pd.DataFrame) -> str:
    lookup = summary.set_index(["dataset_id", "ablation_id"])

    def row(dataset_id: str, ablation_id: str) -> pd.Series:
        return lookup.loc[(dataset_id, ablation_id)]

    baron_default = row("baron_pancreas", REFERENCE_ABLATION)
    baron_batch = row("baron_pancreas", "batch_residualized")
    kang_default = row("kang_ifnb_pbmc", REFERENCE_ABLATION)
    kang_batch = row("kang_ifnb_pbmc", "batch_residualized")
    pbmc_forced = row("pbmc68k_zheng2017", "force_min_embedding_pcs_10")
    pbmc_default = row("pbmc68k_zheng2017", REFERENCE_ABLATION)
    pdac_default = row("pdac_gse263733", REFERENCE_ABLATION)
    pdac_batch = row("pdac_gse263733", "batch_residualized")
    lines = [
        "# Real-data ablation figure and table assets",
        "",
        "Generated by `python scripts/build_realdata_ablation_assets.py`.",
        "",
        "## Evidence Boundary",
        "",
        "- This is a four-dataset `subsample80_pilot10` asset layer: 80% cell subsampling, 10 repeats, five key variants.",
        "- These outputs are suitable for internal/external pre-review and supplemental planning.",
        "- They are not final 20-50 repeat manuscript-grade ablation evidence.",
        "- PBMC68k/Zheng 2017 remains a diagnostic no-call stress case, not a rescued positive discovery.",
        "",
        "## Key Pilot Findings",
        "",
        f"- Baron pancreas: default label ARI {_format_ci(baron_default.mean_label_ari, baron_default.label_ari_ci95_low, baron_default.label_ari_ci95_high)}; batch-residualized label ARI {_format_ci(baron_batch.mean_label_ari, baron_batch.label_ari_ci95_low, baron_batch.label_ari_ci95_high)} and lower batch ARI {_format_ci(baron_batch.mean_batch_ari, baron_batch.batch_ari_ci95_low, baron_batch.batch_ari_ci95_high)} versus default {_format_ci(baron_default.mean_batch_ari, baron_default.batch_ari_ci95_low, baron_default.batch_ari_ci95_high)}.",
        f"- Kang IFN-beta PBMC: default label ARI {_format_ci(kang_default.mean_label_ari, kang_default.label_ari_ci95_low, kang_default.label_ari_ci95_high)}; batch-residualized label ARI {_format_ci(kang_batch.mean_label_ari, kang_batch.label_ari_ci95_low, kang_batch.label_ari_ci95_high)} with lower batch ARI {_format_ci(kang_batch.mean_batch_ari, kang_batch.batch_ari_ci95_low, kang_batch.batch_ari_ci95_high)} versus default {_format_ci(kang_default.mean_batch_ari, kang_default.batch_ari_ci95_low, kang_default.batch_ari_ci95_high)}.",
        f"- PBMC68k/Zheng 2017: default label ARI {_format_ci(pbmc_default.mean_label_ari, pbmc_default.label_ari_ci95_low, pbmc_default.label_ari_ci95_high)} with no-call rate {pbmc_default.no_call_rate:.2f}; forced 10 PCs label ARI {_format_ci(pbmc_forced.mean_label_ari, pbmc_forced.label_ari_ci95_low, pbmc_forced.label_ari_ci95_high)} with no-call rate {pbmc_forced.no_call_rate:.2f}, so this remains a forced-clustering negative control.",
        f"- PDAC GSE263733: default label ARI {_format_ci(pdac_default.mean_label_ari, pdac_default.label_ari_ci95_low, pdac_default.label_ari_ci95_high)}; batch-residualized label ARI {_format_ci(pdac_batch.mean_label_ari, pdac_batch.label_ari_ci95_low, pdac_batch.label_ari_ci95_high)} with lower batch ARI {_format_ci(pdac_batch.mean_batch_ari, pdac_batch.batch_ari_ci95_low, pdac_batch.batch_ari_ci95_high)} versus default {_format_ci(pdac_default.mean_batch_ari, pdac_default.batch_ari_ci95_low, pdac_default.batch_ari_ci95_high)}.",
        "",
        "## Generated Artifacts",
        "",
        f"- Figure source data: `{_rel(SOURCE_DATA)}`",
        f"- Forest plot PNG: `{_rel(FIG_PNG)}`",
        f"- Forest plot PDF: `{_rel(FIG_PDF)}`",
        f"- Forest plot TIFF: `{_rel(FIG_TIFF)}`",
        f"- Figure manifest: `{_rel(FIG_MANIFEST)}`",
        f"- Supplemental table TSV: `{_rel(SUPP_TABLE_TSV)}`",
        f"- Supplemental table DOCX: `{_rel(SUPP_TABLE_DOCX)}`",
        f"- Table manifest: `{_rel(TABLE_MANIFEST)}`",
        "",
        "## Next Use",
        "",
        "Use this asset as a pre-review supplement and as the template for the final 20-50 repeat ablation run. The final manuscript must still add matched Seurat v5, JackStraw or permutation-PCA baselines on the same split framework.",
    ]
    return "\n".join(lines)


def main() -> int:
    summary = build_summary()
    supplemental_table = build_supplemental_table(summary)
    _atomic_write_tsv(summary, SOURCE_DATA)
    _atomic_write_tsv(supplemental_table, SUPP_TABLE_TSV)
    write_docx(supplemental_table)
    render_forest(summary)
    write_manifests()
    _atomic_write_text(build_doc(summary), DOC)
    for path in [
        SOURCE_DATA,
        SUPP_TABLE_TSV,
        SUPP_TABLE_DOCX,
        FIG_PNG,
        FIG_PDF,
        FIG_TIFF,
        FIG_MANIFEST,
        TABLE_MANIFEST,
        DOC,
    ]:
        print(_rel(path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
