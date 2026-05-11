from __future__ import annotations

"""Build a corresponding-author sign-off packet for the Nature Methods route.

Author: RMTGuard development team
Date: 2026-05-10
Purpose: Produce Markdown, TSV, and DOCX handoff files that ask the
corresponding authors to acknowledge the bounded Figure 4 claim before a
Nature Methods presubmission inquiry is sent.
Data source: Generated go/no-go and Figure 4 claim-boundary artifacts.
Method notes: This is an author-confirmation package, not a submission file.
It must preserve the no-guarantee and no-clinical-claim boundaries.
"""

import csv
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]

GO_NO_GO_MD = ROOT / "docs" / "nature_methods_go_no_go_final.md"
FIGURE4_FREEZE_MD = ROOT / "docs" / "figure4_pdac_tme_wording_freeze.md"
FIGURE4_CAPTION_MD = (
    ROOT / "manuscript" / "figure4_caption_strengthened_draft.md"
)
PRESUBMISSION_MD = ROOT / "manuscript" / "nature_methods_presubmission_inquiry.md"
ACK_TEMPLATE_MD = (
    ROOT
    / "manuscript"
    / "corresponding_author_figure4_acknowledgement_template.md"
)

OUT_MD = ROOT / "manuscript" / "corresponding_author_signoff_packet.md"
OUT_TSV = ROOT / "results" / "submission" / "corresponding_author_signoff_packet.tsv"
OUT_DOC_DIR = ROOT / "output" / "doc"
OUT_DOCX = OUT_DOC_DIR / "RMTGuard_corresponding_author_signoff_packet.docx"
OUT_README = OUT_DOC_DIR / "README.md"


CONFIRMATION_TEXT = (
    "I confirm that Figure 4 will be presented only as a bounded public-data "
    "PDAC/TME showcase of RMTGuard callability, not as a new PDAC mechanism, "
    "CAF discovery, prognosis, therapy-response, clinical-validation, or "
    "patient-level claim."
)


def _rel(path: Path) -> str:
    try:
        return str(path.resolve().relative_to(ROOT.resolve())).replace("\\", "/")
    except ValueError:
        return str(path)


def _read_text(path: Path) -> str:
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8", errors="replace")


def _write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(text.rstrip() + "\n", encoding="utf-8")
    tmp.replace(path)


def _write_tsv(path: Path, rows: list[dict[str, str]], fieldnames: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    with tmp.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, delimiter="\t")
        writer.writeheader()
        writer.writerows(rows)
    tmp.replace(path)


def _caption_text() -> str:
    return _read_text(FIGURE4_CAPTION_MD).strip()


def build_rows() -> list[dict[str, str]]:
    return [
        {
            "item_id": "required_decision",
            "status": "author_action_needed",
            "text": CONFIRMATION_TEXT,
            "evidence_path": _rel(FIGURE4_FREEZE_MD),
            "risk_if_unconfirmed": (
                "Nature Methods presubmission inquiry remains blocked."
            ),
        },
        {
            "item_id": "full_submission_status",
            "status": "no_go",
            "text": "Full Nature Methods submission remains no-go.",
            "evidence_path": _rel(GO_NO_GO_MD),
            "risk_if_unconfirmed": (
                "Submitting a full manuscript now would overrun current evidence."
            ),
        },
        {
            "item_id": "presubmission_status",
            "status": "conditional_go_after_author_ack",
            "text": (
                "Nature Methods presubmission inquiry can proceed only after "
                "corresponding-author acknowledgement and final wording review."
            ),
            "evidence_path": _rel(PRESUBMISSION_MD),
            "risk_if_unconfirmed": "Do not send the inquiry.",
        },
        {
            "item_id": "forbidden_claims",
            "status": "hard_stop",
            "text": (
                "Do not describe Figure 4 as a new PDAC mechanism, CAF "
                "discovery, prognosis, therapy-response, clinical-validation, "
                "or patient-level finding."
            ),
            "evidence_path": _rel(FIGURE4_FREEZE_MD),
            "risk_if_unconfirmed": "High desk-reject or reviewer-attack risk.",
        },
    ]


def build_markdown(rows: list[dict[str, str]]) -> str:
    today = date.today().isoformat()
    lines = [
        "# RMTGuard Corresponding Author Sign-off Packet",
        "",
        f"Generated: {today}",
        "",
        "## Mentor Decision",
        "",
        "- Full Nature Methods submission: `NO-GO`.",
        "- Nature Methods presubmission inquiry: `conditional go after author acknowledgement`.",
        "- Genome Biology fallback: `ready if Nature Methods is not encouraging`.",
        "- Acceptance guarantee: `impossible`.",
        "",
        "## Required Confirmation",
        "",
        "> " + CONFIRMATION_TEXT,
        "",
        "## Exact Figure 4 Caption Draft",
        "",
        _caption_text(),
        "",
        "## Sign-off Table",
        "",
        "| Item | Status | Text | Evidence | Risk if not confirmed |",
        "| --- | --- | --- | --- | --- |",
    ]
    for row in rows:
        lines.append(
            "| "
            + row["item_id"]
            + " | `"
            + row["status"]
            + "` | "
            + row["text"]
            + " | `"
            + row["evidence_path"]
            + "` | "
            + row["risk_if_unconfirmed"]
            + " |"
        )
    lines.extend(
        [
            "",
            "## Written Reply Template",
            "",
            "```text",
            "We confirm the bounded Figure 4 wording for RMTGuard.",
            CONFIRMATION_TEXT,
            "",
            "Yi Miao, MD, PhD: confirmed / not confirmed",
            "Han Yan, MD, PhD: confirmed / not confirmed",
            "Date:",
            "Notes:",
            "```",
            "",
            "## Source Files",
            "",
            f"- Go/no-go packet: `{_rel(GO_NO_GO_MD)}`",
            f"- Figure 4 wording freeze: `{_rel(FIGURE4_FREEZE_MD)}`",
            f"- Figure 4 caption draft: `{_rel(FIGURE4_CAPTION_MD)}`",
            f"- Presubmission inquiry draft: `{_rel(PRESUBMISSION_MD)}`",
            f"- Acknowledgement template: `{_rel(ACK_TEMPLATE_MD)}`",
        ]
    )
    return "\n".join(lines)


def _set_document_style(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            styles[style_name].font.name = "Arial"


def _add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def build_docx(rows: list[dict[str, str]]) -> None:
    OUT_DOC_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    _set_document_style(document)
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = document.add_heading("RMTGuard Corresponding Author Sign-off Packet", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(
        "Nature Methods presubmission route control - generated "
        + date.today().isoformat()
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("Mentor Decision", level=1)
    _add_bullet(document, "Full Nature Methods submission: NO-GO.")
    _add_bullet(
        document,
        "Nature Methods presubmission inquiry: conditional go after corresponding-author acknowledgement.",
    )
    _add_bullet(
        document,
        "Genome Biology fallback: ready if Nature Methods presubmission is not encouraging.",
    )
    _add_bullet(document, "Acceptance guarantee: impossible.")

    document.add_heading("Required Confirmation", level=1)
    paragraph = document.add_paragraph()
    run = paragraph.add_run(CONFIRMATION_TEXT)
    run.bold = True

    document.add_heading("Exact Figure 4 Caption Draft", level=1)
    document.add_paragraph(_caption_text())

    document.add_heading("Sign-off Table", level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = [
        "Item",
        "Status",
        "Text",
        "Evidence",
        "Risk if not confirmed",
    ]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row["item_id"]
        cells[1].text = row["status"]
        cells[2].text = row["text"]
        cells[3].text = row["evidence_path"]
        cells[4].text = row["risk_if_unconfirmed"]

    document.add_heading("Signature / Written Confirmation", level=1)
    document.add_paragraph("Yi Miao, MD, PhD: ____________________  Date: __________")
    document.add_paragraph("Han Yan, MD, PhD: ____________________  Date: __________")

    document.add_heading("Written Reply Template", level=1)
    for line in [
        "We confirm the bounded Figure 4 wording for RMTGuard.",
        CONFIRMATION_TEXT,
        "Yi Miao, MD, PhD: confirmed / not confirmed",
        "Han Yan, MD, PhD: confirmed / not confirmed",
        "Date:",
        "Notes:",
    ]:
        document.add_paragraph(line)

    tmp = OUT_DOCX.with_suffix(".docx.tmp")
    document.save(tmp)
    tmp.replace(OUT_DOCX)


def build_readme() -> str:
    return """# RMTGuard Author Sign-off DOCX Outputs

This folder contains the Word handoff file for corresponding-author
acknowledgement before any Nature Methods presubmission inquiry is sent.

## Files

- `RMTGuard_corresponding_author_signoff_packet.docx`

## Boundary

This is an author-confirmation packet. It is not a full manuscript submission
package and does not guarantee journal acceptance.

## Regenerate

```bash
python scripts/build_corresponding_author_signoff_packet.py
```
"""


def main() -> int:
    rows = build_rows()
    _write_tsv(
        OUT_TSV,
        rows,
        [
            "item_id",
            "status",
            "text",
            "evidence_path",
            "risk_if_unconfirmed",
        ],
    )
    _write_text(OUT_MD, build_markdown(rows))
    build_docx(rows)
    _write_text(OUT_README, build_readme())
    print(_rel(OUT_TSV))
    print(_rel(OUT_MD))
    print(_rel(OUT_DOCX))
    print(_rel(OUT_README))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
