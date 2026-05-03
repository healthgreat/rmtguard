#!/usr/bin/env python
"""Build publication-style tables for the RMTGuard manuscript package.

Author: RMTGuard development team
Date: 2026-05-01
Purpose: Convert gate, benchmark, and external-review action data into compact
manuscript tables plus a Word table pack.
Data source: results/gates/*, results/figures/source_data/*, and
results/submission/external_review_action_plan.tsv.
Method notes: This script only changes table presentation. It does not change
benchmark values, gate calls, or manuscript claims.
"""

from __future__ import annotations

import csv
import math
from io import StringIO
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "results" / "tables" / "manuscript"
GATE_REPORT = ROOT / "results" / "gates" / "gate_report.tsv"
PUBLIC_BENCHMARK = (
    ROOT
    / "results"
    / "figures"
    / "source_data"
    / "figure3_public_benchmark_summary.tsv"
)
STABILITY = (
    ROOT
    / "results"
    / "figures"
    / "source_data"
    / "figure3_pbmc3k_stability_summary.tsv"
)
ACTION_PLAN = ROOT / "results" / "submission" / "external_review_action_plan.tsv"
DECISION_MAP = ROOT / "results" / "callability" / "no_call_decision_map.tsv"

TABLE1 = OUT_DIR / "table1_submission_gate_summary.tsv"
TABLE2 = OUT_DIR / "table2_public_benchmark_summary.tsv"
TABLE3 = OUT_DIR / "table3_external_review_action_plan.tsv"
MANIFEST = OUT_DIR / "publication_table_manifest.tsv"
DOCX = OUT_DIR / "rmtguard_publication_tables.docx"

METHOD_LABELS = {
    "rmtguard": "RMTGuard",
    "rmtguard_strict_signal": "RMTGuard strict",
    "scanpy_default_like": "Scanpy-like",
    "fixed_pcs_30": "fixed 30 PCs",
    "fixed_pcs_50": "fixed 50 PCs",
    "elbow_rule": "elbow",
    "parallel_analysis": "parallel analysis",
    "jackstraw_like": "JackStraw-like",
}

DATASET_LABELS = {
    "pbmc3k_10x": "PBMC3k",
    "kang_ifnb_pbmc": "Kang IFN-beta PBMC",
    "baron_pancreas": "Baron pancreas",
    "pbmc68k_zheng2017": "PBMC68k",
}

STATUS_COLORS = {
    "pass": "D9EAD3",
    "fail": "F4CCCC",
    "pending": "D9D2E9",
    "borderline": "FCE5CD",
    "blocked": "F4CCCC",
    "blocked_before_submission": "F4CCCC",
    "blocked before submission": "F4CCCC",
    "implemented_pending_feedback_close": "FFF2CC",
    "implemented pending feedback close": "FFF2CC",
}


def _rel(path: Path) -> str:
    try:
        return str(path.resolve().relative_to(ROOT.resolve())).replace("\\", "/")
    except ValueError:
        return str(path)


def _atomic_write_tsv(df: pd.DataFrame, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    df.to_csv(tmp, sep="\t", index=False, quoting=csv.QUOTE_MINIMAL)
    tmp.replace(path)


def _atomic_write_docx(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    doc.save(tmp)
    tmp.replace(path)


def _read_gate_report(path: Path) -> pd.DataFrame:
    lines = path.read_text(encoding="utf-8").splitlines()
    header_idx = next(i for i, line in enumerate(lines) if line.startswith("gate_id\t"))
    return pd.read_csv(StringIO("\n".join(lines[header_idx:])), sep="\t")


def _fmt(value: object, digits: int = 3) -> str:
    if value is None:
        return "-"
    if isinstance(value, float) and math.isnan(value):
        return "-"
    text = str(value)
    if text.lower() in {"nan", "none", ""}:
        return "-"
    try:
        number = float(text)
    except ValueError:
        return text
    if math.isnan(number):
        return "-"
    return f"{number:.{digits}f}"


def _clean_label(value: object) -> str:
    return str(value).replace("_", " ")


def _sentence_label(value: object) -> str:
    text = _clean_label(value)
    if not text:
        return text
    text = text[0].upper() + text[1:]
    replacements = {
        "pdac tme": "PDAC/TME",
        "doi": "DOI",
        "no call": "no-call",
        "noninferiority": "non-inferiority",
        "pbmc": "PBMC",
    }
    lower_text = text.lower()
    for old, new in replacements.items():
        lower_text = lower_text.replace(old, new)
    return lower_text[0].upper() + lower_text[1:]


def _label_status(value: object) -> str:
    return _clean_label(value)


def _label_method(value: object) -> str:
    return METHOD_LABELS.get(str(value), _clean_label(value))


def _label_dataset(value: object) -> str:
    return DATASET_LABELS.get(str(value), _clean_label(value))


def build_table1() -> pd.DataFrame:
    gates = _read_gate_report(GATE_REPORT)
    gates = gates[gates["gate_id"] != "overall"].copy()
    return pd.DataFrame(
        {
            "Gate": gates["gate_id"].map(_sentence_label),
            "Category": gates["category"].map(_clean_label),
            "Status": gates["status"].map(_label_status),
            "Requirement": gates["nature_methods_requirement"].map(str),
        }
    )


def build_table2() -> pd.DataFrame:
    public = pd.read_csv(PUBLIC_BENCHMARK, sep="\t")
    stability = pd.read_csv(STABILITY, sep="\t")
    decision_map = pd.read_csv(DECISION_MAP, sep="\t")
    decision_map = decision_map[decision_map["unit_type"] == "real_public"][
        ["unit_id", "decision", "primary_reason"]
    ].rename(columns={"unit_id": "dataset_id"})
    keep_methods = [
        "rmtguard",
        "fixed_pcs_30",
        "elbow_rule",
        "scanpy_default_like",
        "jackstraw_like",
    ]
    public = public[public["method"].isin(keep_methods)].copy()
    stability = stability[stability["method"].isin(keep_methods)].copy()
    merged = public.merge(
        stability[
            [
                "dataset_id",
                "method",
                "mean_pairwise_ari",
                "mean_cluster_n",
                "n_repeats",
                "sample_fraction",
            ]
        ],
        on=["dataset_id", "method"],
        how="left",
    )
    merged = merged.merge(decision_map, on="dataset_id", how="left")
    merged["Callability note"] = merged.apply(_callability_note, axis=1)
    merged = merged.sort_values(["dataset_id", "method"])
    return pd.DataFrame(
        {
            "Dataset": merged["dataset_id"].map(_label_dataset),
            "Method": merged["method"].map(_label_method),
            "Annotation ARI": merged["ari"].map(lambda x: _fmt(x, 3)),
            "Annotation NMI": merged["nmi"].map(lambda x: _fmt(x, 3)),
            "Stability ARI": merged["mean_pairwise_ari"].map(lambda x: _fmt(x, 3)),
            "Mean clusters": merged["mean_cluster_n"].map(lambda x: _fmt(x, 1)),
            "Callability note": merged["Callability note"],
        }
    )


def _callability_note(row: pd.Series) -> str:
    if row.get("method") == "rmtguard":
        decision = row.get("decision", "")
        if isinstance(decision, str) and decision.strip():
            return decision.replace("_", " ").replace("no call", "no-call")
    reason = row.get("no_call_reason", "")
    if isinstance(reason, str) and reason.strip():
        return reason.strip().replace("_", " ")
    status = row.get("analysis_status", "")
    if isinstance(status, str) and status.strip() and status.strip() != "ok":
        return status.strip().replace("_", " ")
    if row.get("method") == "rmtguard" and row.get("dataset_id") == "pbmc68k_zheng2017":
        return "diagnostic no-call"
    return "-"


def build_table3() -> pd.DataFrame:
    action = pd.read_csv(ACTION_PLAN, sep="\t")
    return pd.DataFrame(
        {
            "Action": action["action_id"].map(_sentence_label),
            "Priority": action["priority"].map(str),
            "Phase": action["phase"].map(_clean_label),
            "Status": action["status"].map(_label_status),
            "Required action": action["required_action"].map(str),
            "Success gate": action["success_gate"].map(str),
            "Route effect": action["route_effect"].map(str),
        }
    )


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(
    cell, bold: bool = False, size: float = 7.0, color: str | None = None
) -> None:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            run.font.size = Pt(size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = RGBColor.from_string(color)


def _add_dataframe_table(
    doc: Document,
    title: str,
    caption: str,
    df: pd.DataFrame,
    status_column: str | None = None,
) -> None:
    doc.add_heading(title, level=1)
    para = doc.add_paragraph(caption)
    para.style = doc.styles["Caption"]
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0]
    for idx, column in enumerate(df.columns):
        cell = header.cells[idx]
        cell.text = str(column)
        _set_cell_shading(cell, "1F4E79")
        _set_cell_text(cell, bold=True, size=7.2, color="FFFFFF")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, column in enumerate(df.columns):
            value = str(row[column])
            cells[idx].text = value
            _set_cell_text(cells[idx], size=6.6)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if status_column and column == status_column:
                fill = STATUS_COLORS.get(
                    value.lower().replace(" ", "_"),
                    STATUS_COLORS.get(value.lower(), "FFFFFF"),
                )
                _set_cell_shading(cells[idx], fill)
    doc.add_paragraph()


def _build_docx(
    table1: pd.DataFrame, table2: pd.DataFrame, table3: pd.DataFrame
) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RMTGuard Publication Tables")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph(
        "Draft table pack generated from locked source-data tables. Values are not manually edited."
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    warning = doc.add_paragraph(
        "Current submission status: blocked before journal submission until public release/DOI and remaining external-review action gates are resolved."
    )
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning.runs[0].font.color.rgb = RGBColor(178, 24, 43)
    warning.runs[0].font.bold = True

    _add_dataframe_table(
        doc,
        "Table 1. Submission Gate Summary",
        "Evidence-gated readiness checks. This table preserves the current fail/pending calls rather than masking them.",
        table1,
        status_column="Status",
    )
    _add_dataframe_table(
        doc,
        "Table 2. Public Benchmark Summary",
        "Public benchmark values used for Figure 3. Missing annotation entries are shown as '-'.",
        table2,
    )
    _add_dataframe_table(
        doc,
        "Table 3. External Review Action Plan",
        "Prioritized response plan derived from the active external-review feedback triage.",
        table3,
        status_column="Status",
    )
    return doc


def _write_manifest(paths: Iterable[Path]) -> None:
    rows = [
        {
            "artifact": path.name,
            "path": _rel(path),
            "status": "written" if path == MANIFEST or path.exists() else "missing",
            "regeneration_command": "python scripts/build_publication_tables.py",
        }
        for path in paths
    ]
    _atomic_write_tsv(pd.DataFrame(rows), MANIFEST)


def main() -> None:
    table1 = build_table1()
    table2 = build_table2()
    table3 = build_table3()
    _atomic_write_tsv(table1, TABLE1)
    _atomic_write_tsv(table2, TABLE2)
    _atomic_write_tsv(table3, TABLE3)
    doc = _build_docx(table1, table2, table3)
    _atomic_write_docx(doc, DOCX)
    _write_manifest([TABLE1, TABLE2, TABLE3, DOCX, MANIFEST])
    print(_rel(MANIFEST))


if __name__ == "__main__":
    main()
