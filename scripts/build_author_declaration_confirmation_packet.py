"""Build an author-declaration confirmation packet for RMTGuard.

Author: RMTGuard contributors
Date: 2026-05-12
Purpose: Convert manuscript author declarations, CRediT roles, funding,
competing interests, ethics/public-data wording, and Figure 4 acknowledgement
blockers into an auditable confirmation packet before any v0.1.1 release or
editor-facing submission.
Data source: metadata/author_metadata.tsv, metadata/credit_roles.tsv,
metadata/corresponding_author_signoff_tracker.tsv, and
results/submission/reporting_summary_draft.tsv.
Method notes: This script does not certify author approval. It only creates a
checklist and Word packet for manual author confirmation.
"""

from __future__ import annotations

import csv
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
AUTHOR_META = ROOT / "metadata" / "author_metadata.tsv"
CREDIT_ROLES = ROOT / "metadata" / "credit_roles.tsv"
SIGNOFF = ROOT / "metadata" / "corresponding_author_signoff_tracker.tsv"
REPORTING = ROOT / "results" / "submission" / "reporting_summary_draft.tsv"

OUT_TSV = ROOT / "results" / "submission" / "author_declaration_confirmation_checklist.tsv"
OUT_MD = ROOT / "docs" / "author_declaration_confirmation_packet.md"
OUT_DOCX = ROOT / "output" / "doc" / "RMTGuard_author_declaration_confirmation_packet.docx"


@dataclass(frozen=True)
class ConfirmationItem:
    item_id: str
    category: str
    owner: str
    current_value: str
    current_status: str
    confirmation_question: str
    required_author_action: str
    evidence_to_save: str
    gate_unblocked: str
    stop_condition: str


FIELDNAMES = [
    "item_id",
    "category",
    "owner",
    "current_value",
    "current_status",
    "confirmation_question",
    "required_author_action",
    "evidence_to_save",
    "gate_unblocked",
    "stop_condition",
]


def _rel(path: Path) -> str:
    try:
        return path.resolve().relative_to(ROOT.resolve()).as_posix()
    except ValueError:
        return str(path)


def read_tsv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle, delimiter="\t"))


def write_tsv(path: Path, rows: list[ConfirmationItem]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    with tmp.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=FIELDNAMES, delimiter="\t")
        writer.writeheader()
        for row in rows:
            writer.writerow(row.__dict__)
    tmp.replace(path)


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(text.rstrip() + "\n", encoding="utf-8")
    tmp.replace(path)


def meta_by_field(rows: list[dict[str, str]]) -> dict[str, dict[str, str]]:
    return {row.get("field", ""): row for row in rows}


def count_pending_credit_rows(rows: list[dict[str, str]]) -> int:
    return sum(1 for row in rows if row.get("status") != "author_confirmed")


def credit_summary(rows: list[dict[str, str]]) -> str:
    by_author: dict[str, list[str]] = {}
    for row in rows:
        author = row.get("author", "unknown")
        role = row.get("credit_role", "")
        level = row.get("contribution_level", "")
        by_author.setdefault(author, []).append(f"{role} ({level})")
    return " | ".join(f"{author}: {', '.join(roles)}" for author, roles in by_author.items())


def reporting_summary(rows: list[dict[str, str]]) -> str:
    pending = [
        f"{row.get('section')}/{row.get('item')}"
        for row in rows
        if row.get("status") in {"pending_manual", "needs_author_completion"}
    ]
    return "; ".join(pending) if pending else "No pending manual rows."


def signoff_summary(rows: list[dict[str, str]]) -> str:
    if not rows:
        return "No corresponding-author signoff tracker rows."
    return "; ".join(
        f"{row.get('author_name')}={row.get('status')}" for row in rows
    )


def build_items() -> list[ConfirmationItem]:
    author_rows = read_tsv(AUTHOR_META)
    credit_rows = read_tsv(CREDIT_ROLES)
    signoff_rows = read_tsv(SIGNOFF)
    reporting_rows = read_tsv(REPORTING)
    meta = meta_by_field(author_rows)

    postal_value = "author_provided=350000; public_candidate=210019"
    if "postal_code_author_provided" in meta or "postal_code_public_source_candidate" in meta:
        postal_value = (
            f"author_provided={meta.get('postal_code_author_provided', {}).get('value', '')}; "
            f"public_candidate={meta.get('postal_code_public_source_candidate', {}).get('value', '')}"
        )

    return [
        ConfirmationItem(
            "postal_code_final",
            "title_page",
            "Chongfa Chen and corresponding authors",
            postal_value,
            meta.get("postal_code_author_provided", {}).get("status", "missing"),
            "Which postal code should appear in the title page and submission metadata?",
            "Confirm the final postal code in writing; if 350000 is correct, explicitly state that it should override the public-source candidate.",
            "metadata/author_reply_evidence/postal_code_confirmation.*",
            "author_declarations",
            "Do not finalize title page, .zenodo.json, or submission metadata while this conflict remains unresolved.",
        ),
        ConfirmationItem(
            "credit_roles_all_authors",
            "author_contributions",
            "All authors",
            credit_summary(credit_rows),
            f"{count_pending_credit_rows(credit_rows)} rows pending author confirmation",
            "Do all authors approve these CRediT roles and contribution levels?",
            "Each author should approve or edit their own role list before submission.",
            "metadata/author_reply_evidence/credit_roles_confirmation.*",
            "author_declarations",
            "Do not paste CRediT roles into the journal system until all authors confirm.",
        ),
        ConfirmationItem(
            "funding_statement_final",
            "funding",
            "Corresponding authors",
            meta.get("funding_statement", {}).get("value", "missing"),
            meta.get("funding_statement", {}).get("status", "missing"),
            "Was this work supported by specific grants or should we state that no specific funding was received?",
            "Provide exact funder names and grant numbers, or confirm the no-specific-funding wording.",
            "metadata/author_reply_evidence/funding_statement_confirmation.*",
            "author_declarations",
            "Do not submit or mint v0.1.1 metadata with a placeholder funding statement.",
        ),
        ConfirmationItem(
            "competing_interests_final",
            "competing_interests",
            "All authors",
            meta.get("competing_interests_statement", {}).get("value", "missing"),
            meta.get("competing_interests_statement", {}).get("status", "missing"),
            "Can all authors confirm that there are no financial or non-financial competing interests?",
            "All authors should confirm or disclose conflicts before the statement is used.",
            "metadata/author_reply_evidence/competing_interests_confirmation.*",
            "author_declarations",
            "Do not use 'no competing interests' unless every author confirms it.",
        ),
        ConfirmationItem(
            "ethics_public_data_final",
            "ethics",
            "Corresponding authors",
            meta.get("ethics_public_data_statement", {}).get("value", "missing"),
            meta.get("ethics_public_data_statement", {}).get("status", "missing"),
            "Can the corresponding authors confirm that the manuscript uses only public, de-identified datasets and no new private clinical data?",
            "Confirm the final dataset list is public-data-only and that original studies reported consent/ethics where applicable.",
            "metadata/author_reply_evidence/ethics_public_data_confirmation.*",
            "author_declarations; reporting_summary",
            "Stop if any private clinical data, unpublished patient metadata, or identifiable information is added.",
        ),
        ConfirmationItem(
            "figure4_bounded_wording_ack",
            "figure_wording",
            "Yi Miao and Han Yan",
            signoff_summary(signoff_rows),
            "proxy_authorized_working_assumption",
            "Do the corresponding authors approve the bounded Figure 4 wording as a public-data use case with no mechanism or clinical claim?",
            "Save written approval and record it with scripts/record_corresponding_author_signoff.py.",
            "metadata/author_reply_evidence/figure4_acknowledgement_*.txt",
            "corresponding_author_figure4_ack; figure_caption_source_audit",
            "Do not send a Nature Methods presubmission inquiry until both corresponding authors are recorded as confirmed.",
        ),
        ConfirmationItem(
            "reporting_summary_author_check",
            "reporting_summary",
            "Corresponding authors",
            reporting_summary(reporting_rows),
            "pending_author_verification",
            "Can the corresponding authors verify the reporting-summary worksheet entries before official submission?",
            "Check statistics, multiple testing, software, data availability, ethics, and AI-use entries against the final manuscript.",
            "metadata/author_reply_evidence/reporting_summary_confirmation.*",
            "reporting_summary",
            "Do not transfer draft answers into the official Nature form without this verification.",
        ),
        ConfirmationItem(
            "title_page_author_metadata_check",
            "title_page",
            "All authors",
            "Chongfa Chen; Han Yan; Yi Miao",
            "author_provided_with_pending_checks",
            "Do all authors approve the author order, affiliations, emails, and ORCID records?",
            "Confirm final author order, corresponding-author order, institutional name, emails, and ORCID IDs.",
            "metadata/author_reply_evidence/title_page_author_metadata_confirmation.*",
            "author_declarations",
            "Stop if any author order, affiliation, email, or ORCID correction is requested.",
        ),
    ]


def build_markdown(items: list[ConfirmationItem]) -> str:
    author_rows = read_tsv(AUTHOR_META)
    credit_rows = read_tsv(CREDIT_ROLES)
    blocked = [row for row in items if "Do not" in row.stop_condition or "Stop" in row.stop_condition]
    lines = [
        "# RMTGuard author declaration confirmation packet",
        "",
        f"Generated: {date.today().isoformat()}",
        "Generated by `python scripts/build_author_declaration_confirmation_packet.py`.",
        "",
        "## Purpose",
        "",
        "This packet converts the remaining author-controlled manuscript blockers into exact confirmation questions. It is not a substitute for author approval and should not be treated as a journal-ready declaration until written confirmations are saved.",
        "",
        "## Current Decision",
        "",
        "- Status: `manual_confirmation_required`.",
        f"- Checklist: `{_rel(OUT_TSV)}`",
        f"- Word packet: `{_rel(OUT_DOCX)}`",
        f"- Confirmation items: `{len(items)}`",
        f"- Stop-rule items: `{len(blocked)}`",
        "",
        "## Confirmation Checklist",
        "",
        "| Item | Owner | Current status | Required action | Gate unblocked |",
        "| --- | --- | --- | --- | --- |",
    ]
    for item in items:
        lines.append(
            f"| {item.item_id} | {item.owner} | {item.current_status} | {item.required_author_action} | {item.gate_unblocked} |"
        )

    lines.extend(
        [
            "",
            "## Exact Confirmation Text To Send",
            "",
            "Authors can reply by editing the text below or by sending written confirmation. Save replies under `metadata/author_reply_evidence/`.",
            "",
        ]
    )
    for item in items:
        lines.extend(
            [
                f"### {item.item_id}",
                "",
                f"Owner: {item.owner}",
                "",
                f"Question: {item.confirmation_question}",
                "",
                f"Current value: `{item.current_value}`",
                "",
                "Confirmation response:",
                "",
                "```text",
                "I confirm this item as written.",
                "OR",
                "Please revise this item as follows: [author edits here].",
                "```",
                "",
                f"Evidence to save: `{item.evidence_to_save}`",
                "",
            ]
        )

    lines.extend(
        [
            "## Draft CRediT Roles",
            "",
            "| Author | Role | Level | Status |",
            "| --- | --- | --- | --- |",
        ]
    )
    for row in credit_rows:
        lines.append(
            f"| {row.get('author', '')} | {row.get('credit_role', '')} | {row.get('contribution_level', '')} | {row.get('status', '')} |"
        )

    lines.extend(
        [
            "",
            "## Author Metadata Rows",
            "",
            "| Field | Value | Status | Required action |",
            "| --- | --- | --- | --- |",
        ]
    )
    for row in author_rows:
        value = row.get("value", "").replace("|", "/")
        action = row.get("required_action", "").replace("|", "/")
        lines.append(
            f"| {row.get('field', '')} | {value} | {row.get('status', '')} | {action} |"
        )

    lines.extend(
        [
            "",
            "## After Authors Reply",
            "",
            "1. Save written replies under `metadata/author_reply_evidence/`.",
            "2. Record Figure 4 replies with `python scripts/record_corresponding_author_signoff.py`.",
            "3. Update `metadata/author_metadata.tsv` and `metadata/credit_roles.tsv` only after confirmation.",
            "4. Re-run `python scripts/build_author_declaration_confirmation_packet.py`.",
            "5. Re-run `python scripts/build_v0_1_1_release_preflight.py`.",
            "",
            "## Boundary",
            "",
            "This packet reduces ambiguity, but it does not create author approval, ethics approval, a GitHub Release, or a Zenodo DOI.",
        ]
    )
    return "\n".join(lines)


def configure_doc(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    for style_name, size in [("Normal", 10), ("Title", 17), ("Heading 1", 14), ("Heading 2", 11)]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        if style_name.startswith("Heading"):
            style.font.bold = True


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        header_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def build_docx(items: list[ConfirmationItem]) -> None:
    credit_rows = read_tsv(CREDIT_ROLES)
    document = Document()
    configure_doc(document)
    document.add_heading("RMTGuard Author Declaration Confirmation Packet", level=0)
    document.add_paragraph(f"Generated: {date.today().isoformat()}")
    document.add_paragraph(
        "Purpose: confirm author-controlled manuscript declarations before any v0.1.1 release or editor-facing submission."
    )
    document.add_heading("Decision", level=1)
    document.add_paragraph(
        "Current status: manual_confirmation_required. This document is a confirmation packet, not author approval."
    )
    document.add_heading("Checklist", level=1)
    add_table(
        document,
        ["Item", "Owner", "Status", "Required action", "Gate"],
        [
            [
                item.item_id,
                item.owner,
                item.current_status,
                item.required_author_action,
                item.gate_unblocked,
            ]
            for item in items
        ],
    )
    document.add_heading("Confirmation Text", level=1)
    for item in items:
        document.add_heading(item.item_id, level=2)
        document.add_paragraph(f"Owner: {item.owner}")
        document.add_paragraph(f"Question: {item.confirmation_question}")
        document.add_paragraph(f"Current value: {item.current_value}")
        document.add_paragraph("Response: I confirm this item as written. / Please revise this item as follows: ____")
        document.add_paragraph(f"Evidence to save: {item.evidence_to_save}")

    document.add_heading("Draft CRediT Roles", level=1)
    add_table(
        document,
        ["Author", "Role", "Level", "Status"],
        [
            [
                row.get("author", ""),
                row.get("credit_role", ""),
                row.get("contribution_level", ""),
                row.get("status", ""),
            ]
            for row in credit_rows
        ],
    )
    document.add_heading("Boundary", level=1)
    document.add_paragraph(
        "Do not use this document as journal-ready approval until written replies are saved and the release preflight is re-run."
    )
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    tmp = OUT_DOCX.with_suffix(OUT_DOCX.suffix + ".tmp")
    document.save(tmp)
    tmp.replace(OUT_DOCX)


def main() -> int:
    items = build_items()
    write_tsv(OUT_TSV, items)
    write_text(OUT_MD, build_markdown(items))
    build_docx(items)
    print(_rel(OUT_TSV))
    print(_rel(OUT_MD))
    print(_rel(OUT_DOCX))
    print("manual_confirmation_required")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
